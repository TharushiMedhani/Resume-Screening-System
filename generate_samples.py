from docx import Document
import os

def create_resume(filename, name, role, skills, experience):
    doc = Document()
    doc.add_heading(name, 0)
    doc.add_paragraph(role)
    
    doc.add_heading('Professional Summary', level=1)
    doc.add_paragraph(f"Dedicated {role} with extensive experience in software development.")
    
    doc.add_heading('Technical Skills', level=1)
    doc.add_paragraph(", ".join(skills))
    
    doc.add_heading('Experience', level=1)
    for exp in experience:
        doc.add_heading(exp['title'], level=2)
        doc.add_paragraph(exp['description'])
    
    doc_path = os.path.join('samples', filename)
    os.makedirs('samples', exist_ok=True)
    doc.save(doc_path)
    print(f"Created {doc_path}")

# Strong Resume (AI/DS Focus)
create_resume(
    'Strong_AI_Resume.docx',
    'John Expert',
    'Senior AI Engineer',
    ['Python', 'PyTorch', 'TensorFlow', 'Machine Learning', 'Deep Learning', 'FastAPI', 'Docker', 'AWS', 'SQL', 'NLP', 'LLM', 'Generative AI'],
    [
        {'title': 'Lead AI Engineer at Google', 'description': 'Led the development of large scale LLM models using Generative AI and PyTorch. Built APIs using FastAPI and deployed on AWS with Docker.'},
        {'title': 'Data Scientist at Meta', 'description': 'Implemented deep learning algorithms for NLP tasks. Used SQL for data analysis and visualization.'}
    ]
)

# Average Resume (Web Dev Focus)
create_resume(
    'Average_Web_Resume.docx',
    'Jane Mid',
    'Full Stack Developer',
    ['JavaScript', 'React', 'Node.js', 'Express', 'HTML', 'CSS', 'Git', 'MySQL', 'Unit Testing'],
    [
        {'title': 'Web Developer at StartUp Inc.', 'description': 'Built responsive web designs using React and CSS. Developed backend services with Node.js and Express.'},
        {'title': 'Junior Dev at Agency', 'description': 'Developed landing pages using HTML and CSS. Used Git for version control.'}
    ]
)

# Weak Resume (Finance/General Focus)
create_resume(
    'Weak_General_Resume.docx',
    'Bob Novice',
    'Office Administrator',
    ['Communication', 'Leadership', 'Time Management', 'Microsoft Excel', 'Customer Service'],
    [
        {'title': 'Admin Assistant', 'description': 'Managed office supplies and coordinated team meetings. Provided excellent customer service and used Excel for reporting.'},
        {'title': 'Customer Support Representative', 'description': 'Handled customer inquiries and resolved issues in a timely manner. Developed leadership skills.'}
    ]
)
